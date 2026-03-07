from openpyxl import Workbook
from docx import Document
from docx.shared import Pt
from datetime import datetime
import re
import io

PASS_MARK = 33

def process_result(uploaded_file, subject_inputs=None, teacher_inputs=None):
    """
    Direct conversion of your working script to work with Flask:
    - uploaded_file: file-like object (Flask's FileStorage or BytesIO)
    - subject_inputs: optional dict to override/add SUBJECT_MAP entries
    - teacher_inputs: optional dict to override/add teacher_map entries

    Returns:
      - {"missing_subjects": [...]}  if new subject codes need mapping
      - {"missing_teachers": [...]}  if teacher names are missing
      - {"excel_file": BytesIO, "word_file": BytesIO} on success
    """

    # === base mappings (same as your working script) ===
    SUBJECT_MAP = {
        "184": "English",
        "301":"English",
        "302":"Hindi core",
        "028":"Political Science",
        "030":"Economics",
        "037":"Psychology",
        "042":"Physics",
        "043":" Chemistry ",
        "044":"Biology",
        "048":"Physical Education",
        "049":"Painting",
        "283":"CS old",
        "002": "Hindi-A",
        "085": "Hindi-B",
        "122":"Sanskrit",
        "041": "Maths Standard",
        "241": "Maths Basic",
        "086": "Science",
        "087": "SST",
        "402": "IT",
        "055":"Accountancy",
        "054":"Business Studies",
    }

    teacher_map = {
    }

    # Allow overrides from the web UI (so UI can resubmit with mappings)
    if subject_inputs:
        SUBJECT_MAP.update(subject_inputs)
    if teacher_inputs:
        teacher_map.update(teacher_inputs)

    # === read uploaded file content ===
    content = uploaded_file.read().decode("utf-8", errors="ignore")

    # preserve your exact trimming logic
    lines = []
    for l in content.split("\n"):
        if l.strip():
            lines.append(l.rstrip())

    students = {}
    all_subjects = set()
    missing_subject_codes = set()

    # helper to detect result token in the surrounding lines (non-destructive)
    def detect_result_token(check_lines):
        # returns normalized token or None
        # check for RL, ABST, COMP, PASS, ER/FAIL, WITHHELD
        for s in check_lines:
            if not s:
                continue
            # common RL forms
            if re.search(r'\bR\.?L\.?\b', s, re.I) or re.search(r'\bRESULT\s+LATER\b', s, re.I) or re.search(r'\bWITHHELD\b', s, re.I):
                return "R.L."
            if re.search(r'\bABST\b', s, re.I) or re.search(r'\bAB\b', s, re.I) or re.search(r'\bABSENT\b', s, re.I):
                return "ABSENT"
            if re.search(r'\bCOMP\b', s, re.I) or re.search(r'\bCOMPARTMENT\b', s, re.I):
                return "COMP"
            if re.search(r'\bPASS\b', s, re.I):
                return "PASS"
            if re.search(r'\bER\b', s, re.I) or re.search(r'\bFAIL\b', s, re.I):
                return "FAIL"
        return None

    i = 0
    while i < len(lines):
        line = lines[i]

        # Detect a student line by roll no = 8 digit numeric
        if re.match(r"^\d{8}", line):
            parts = line.split()
            roll = parts[0]                      # roll number
            gender = parts[1] if len(parts) > 1 else ""   # M/F if present
            name = ''

            n = 2
            # same name-collection logic as your working script
            while n < len(parts) and not parts[n].isdigit():
                name += parts[n] + ' '
                n += 1

            name = name.strip()

            subject_codes = []
            while n < len(parts) and re.fullmatch(r"\d{3}", parts[n]):
                subject_codes.append(parts[n])
                n += 1

            # Convert subject codes to readable names (keep code placeholder if unknown)
            subject_names = []
            for code in subject_codes:
                if code in SUBJECT_MAP:
                    subject_names.append(SUBJECT_MAP[code])
                else:
                    # keep code as placeholder (this mirrors your working code)
                    subject_names.append(code)
                    # record missing code so UI can ask for mapping
                    missing_subject_codes.add(code)

            # Move to possible marks/grade line
            i += 1
            marks_line = []

            while i < len(lines):
                next_line = lines[i].strip()

                # Stop if this is a new roll number → means no marks line for current student
                if re.match(r"^\d{8}", next_line):
                    break

                # If it's a marks line (starts with marks like "072 B1 ...")
                if re.match(r"^\d{2,3}\s", next_line):
                    marks_line = next_line.split()
                    break

                i += 1

            # --- Extract marks & grades similar to your script ---
            marks_list = []
            grades_list = []

            # Extract marks & grades in pairs
            for j in range(0, len(marks_line), 2):
                if j + 1 >= len(marks_line):
                    break

                mark = marks_line[j]
                grade = marks_line[j+1]

                if mark.isdigit():
                    marks_list.append(int(mark))
                    grades_list.append(grade)
            else:
                # this 'else' mirrors your working script (runs if for-loop completes without break)
                # in your script you used an else to handle missing marks_line — keep behavior identical
                # (practically this branch usually won't run if marks_line exists)
                marks_list = [0] * len(subject_codes)
                i -= 1

            # In your working script you then re-extracted marks_list only (keeping same behavior)
            marks_list = []
            for j in range(0, len(marks_line), 2):
                if j + 1 >= len(marks_line):
                    break
                mark = marks_line[j]
                if mark.isdigit():
                    marks_list.append(int(mark))

            # Map marks to subjects in fixed CBSE order 
            # (your working script used only codes present in SUBJECT_MAP for mapping)
            mapped_subjects = [SUBJECT_MAP.get(code, code) for code in subject_codes if code in SUBJECT_MAP]

            marks = dict(zip(mapped_subjects, marks_list[:len(mapped_subjects)]))
            grades = dict(zip(mapped_subjects, grades_list[:len(mapped_subjects)]))

            # Track subjects for dynamic Excel creation
            for s in marks.keys():
                all_subjects.add(s)

            # ---- CALCULATE BOTH TOTALS (Main5 / Best5) (kept exactly) ----
            main5_subjects = list(marks.keys())[:5]
            main5_marks = [marks[s] for s in main5_subjects]
            main5_total = sum(main5_marks)
            main5_percent = round(main5_total / len(main5_marks), 2) if main5_marks else 0

            english = marks.get("English", 0)
            other_sub_marks = [m for s, m in marks.items() if s != "English"]
            top4 = sorted(other_sub_marks, reverse=True)[:4]
            best5_total = english + sum(top4)
            best5_percent = round(best5_total / 5, 2)

            # --- Detect RESULT token from the student block (original line + a few following lines)
            # We'll check the current student line and up to two following lines for result keywords
            check_block = [line]
            # i currently points to the line after the student line (or the marks line). Grab next up-to-2 lines safely.
            for k in range(i, min(i+3, len(lines))):
                check_block.append(lines[k])

            result_token = detect_result_token(check_block)  # may be None

            # Store student record (structure identical) plus Result token
            students[roll] = {
                "Gender": gender,
                "Name": name,
                "Marks": marks,
                "Grades": grades,
                "Main5_Total": main5_total,
                "Main5_Percent": main5_percent,
                "Best5_Total": best5_total,
                "Best5_Percent": best5_percent,
                "Result": result_token
            }

        i += 1

    # If any unknown subject codes found, return them so UI can ask (Flask flow)
    if missing_subject_codes:
        return {"missing_subjects": sorted(list(missing_subject_codes))}

    # Detect missing teachers for subjects encountered — return for UI to ask (keeps flow consistent)
    missing_teachers = [s for s in all_subjects if s not in teacher_map]
    if missing_teachers:
        return {"missing_teachers": sorted(missing_teachers)}

    # ================= EXCEL (same as working script) =================
    wb = Workbook()
    sheet = wb.active
    sheet.title = "CBSE Result"

    all_subjects = sorted(all_subjects)
    header = ["Roll No", "Name"] + all_subjects + ["Main5 Total", "Main5 Percentage","Best5 Total","Best5 Percentage"]
    sheet.append(header)

    for roll, data in students.items():
        row = [roll, data["Name"]]
        for sub in all_subjects:
            row.append(data["Marks"].get(sub, ""))
        row.append(data["Main5_Total"])
        row.append(data["Main5_Percent"])
        row.append(data["Best5_Total"])
        row.append(data["Best5_Percent"])
        sheet.append(row)

    # ================= WORD (Forms) and ANALYSIS (kept same) =================
    doc = Document()

    # ===== PAGE 1 — FORM C : COMPARTMENT STUDENTS =====
    doc.add_heading("Form C", level=1).runs[0].font.size = Pt(22)
    doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("Name of the School : SPS TDSC")
    doc.add_paragraph("Details about students placed in compartment")
    doc.add_paragraph("(Class X Session 2024-25)")
    doc.add_paragraph("")

    failed_students = []
    for roll, info in students.items():
        marks = info["Marks"]
        grades = info["Grades"]

        # detect ABSENT using Result token if present, else fallback to marks-based
        result_tok = info.get("Result")
        if result_tok == "ABSENT":
            continue   # ← DO NOT include absent students in compartment list
        if result_tok == "R.L.":
            continue   # ← DO NOT include RL in compartment list (RL should be excluded)

        # fallback: keep original logic if no explicit token
        is_absent = all(m == 0 for m in marks.values())
        if is_absent:
            continue   # ← DO NOT include absent students in compartment list

        failed_subjects = [s for s, m in marks.items() if m < PASS_MARK]

        # COMPARTMENT ONLY IF 1 or 2 failed subjects
        if 1 <= len(failed_subjects) <= 2:
            failed_students.append((info["Name"], failed_subjects))

    # --- Create Form C table ---
    table_c = doc.add_table(rows=1, cols=3)
    table_c.style = "Table Grid"
    hdr = table_c.rows[0].cells
    hdr[0].text = "S. No."
    hdr[1].text = "Name of Student"
    hdr[2].text = "Subject(s) in which placed under Compartment"

    if not failed_students:
        row = table_c.add_row().cells
        row[0].text = "-"
        row[1].text = "No students placed under Compartment"
        row[2].text = "-"
    else:
        for idx, (name, subjects) in enumerate(failed_students, start=1):
            row = table_c.add_row().cells
            row[0].text = str(idx)
            row[1].text = name
            row[2].text = ", ".join(subjects) if subjects else "All Subjects (Absent)"

    for row in table_c.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("Signature of School Principal".ljust(60) + "Signature of Exam I/C")

    # ===== PAGE 2 — FORM D : TOP 25 STUDENTS =====
    doc.add_page_break()
    doc.add_heading("Form D", level=1).runs[0].font.size = Pt(22)
    doc.add_paragraph("Name of the School : SPS TDSC")
    doc.add_paragraph("Details about TOP 25 students/rankers of class X (Session 2024-25)")
    doc.add_paragraph("AGGREGATE MARKS OUT OF 500 (Out of Main Subjects only)")
    doc.add_paragraph("")

    # ---- FILTER ONLY PASS STUDENTS ----
    eligible_students = []
    for roll, info in students.items():
        marks = info["Marks"]
        # determine absence/RL from Result token first, fallback to marks-based
        result_tok = info.get("Result")
        if result_tok == "ABSENT":
            continue
        if result_tok == "R.L.":
            continue

        is_absent = all(m == 0 for m in marks.values()) if marks else True
        if is_absent:
            continue
        main5_subjects = list(marks.keys())[:5]
        main5_marks = [marks[s] for s in main5_subjects]
        fail_count = sum(1 for m in main5_marks if m < PASS_MARK)
        if fail_count == 0:
            eligible_students.append((roll, info))

    sorted_students = sorted(eligible_students, key=lambda x: x[1]["Main5_Total"], reverse=True)
    top25 = sorted_students[:25]
    maximum_marks = max([info["Main5_Total"] for roll, info in eligible_students], default=0)

    table_d = doc.add_table(rows=1, cols=4)
    table_d.style = "Table Grid"
    hdr2 = table_d.rows[0].cells
    hdr2[0].text = "RANK"
    hdr2[1].text = "Name of the Student"
    hdr2[2].text = "Aggregate Marks (out of 500)"
    hdr2[3].text = "%"

    rank = 0
    previous_total = None
    for idx, (roll, info) in enumerate(top25, start=1):
        current_total = info["Main5_Total"]
        if current_total != previous_total:
            rank +=1
        previous_total = current_total
        row = table_d.add_row().cells
        row[0].text = str(rank)
        row[1].text = info["Name"]
        row[2].text = str(info["Main5_Total"])
        row[3].text = str(info["Main5_Percent"])

    for row in table_d.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("Signature of School Principal".ljust(60) + "Signature of EXAM INCHARGE")
    doc.add_page_break()

    # ===== PAGE 3 — FORMAT A : CONSOLIDATED ANALYSIS =====
    doc.add_heading("FORMAT A : CONSOLIDATED ANALYSIS OF SUBJECT-WISE AGGREGATE %", level=1).runs[0].font.size = Pt(16)
    doc.add_paragraph("Name of School : SPS TDSC")
    doc.add_paragraph("CBSE Class X RESULT - Session 2024-25")
    doc.add_paragraph("")

    total_students = len(students)

    present_students = 0
    absent_students = 0
    rl_students = 0
    failed_students = 0
    compartment_students = 0
    pass_students = 0

    grade_count = {"A1":0,"A2":0,"B1":0,"B2":0,"C1":0,"C2":0,"D1":0,"D2":0,"E":0}

    total_main5_sum = 0
    valid_main5_students = 0
    highest_percent = 0
    A1_inall_sub = 0

    for roll, info in students.items():

        marks = info["Marks"]
        grades_list_for_main5 = list(info["Grades"].values())[:5]

        result_token = info.get("Result")

        # ===== RESULT PRIORITY =====

        if result_token == "ABSENT":
            absent_students += 1
            continue

        if result_token == "R.L.":
            rl_students += 1
            present_students += 1
            continue

        # ===== MARKS BASED FALLBACK =====

        is_absent = all(m == 0 for m in marks.values())

        if is_absent:
            absent_students += 1
            continue

        present_students += 1

        # ===== GRADE DISTRIBUTION =====

        for g in grades_list_for_main5:
            if g in grade_count:
                grade_count[g] += 1

        if grades_list_for_main5.count("A1") == 5:
            A1_inall_sub += 1

        # ===== MAIN 5 CALCULATIONS =====

        p = info["Main5_Percent"]

        if p > 0:
            total_main5_sum += p
            valid_main5_students += 1
            highest_percent = max(highest_percent, p)

        main5_subjects = list(info["Marks"].keys())[:5]
        main5_marks = [info["Marks"][x] for x in main5_subjects]

        fail_count = sum(1 for m in main5_marks if m < PASS_MARK)

        if fail_count == 0:
            pass_students += 1

        elif fail_count in (1,2):
            compartment_students += 1

        else:
            failed_students += 1


    # ===== SCHOOL AVERAGE =====

    if valid_main5_students > 0:
        school_avg = round(total_main5_sum / valid_main5_students, 2)
    else:
        school_avg = 0


    # ===== EVALUATED STUDENTS =====

    evaluated_students = present_students - rl_students


    # ===== PASS / FAIL / COMPARTMENT % =====

    if evaluated_students > 0:

        pass_percent = round((pass_students / evaluated_students) * 100, 2)

        fail_percent = round((failed_students / evaluated_students) * 100, 2)

        compartment_percent = round((compartment_students / evaluated_students) * 100, 2)

    else:

        pass_percent = 0
        fail_percent = 0
        compartment_percent = 0


    # ===== SCHOOL SCORE =====

    total_grade_slots = evaluated_students * 5 if evaluated_students > 0 else 1

    school_score = round(
    (
    9*grade_count["A1"] + 8*grade_count["A2"] +
    7*grade_count["B1"] + 6*grade_count["B2"] +
    5*grade_count["C1"] + 4*grade_count["C2"] +
    3*grade_count["D1"] + 2*grade_count["D2"] +
    1*grade_count["E"]
    ) / (9 * total_grade_slots) * 100, 2)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "S. No."
    hdr[1].text = "DETAILS"
    hdr[2].text = "No."
    hdr[3].text = "%"

    rows_data = [
        ("1", "Total Students Appeared", present_students, round(present_students/(present_students+absent_students)*100,2) if (present_students+absent_students)>0 else 0),
        ("2", "Absent Students", absent_students, round(absent_students/(present_students+absent_students)*100,2) if (present_students+absent_students)>0 else 0),
        ("3", "Passed", pass_students, pass_percent),
        ("4", "Compartment", compartment_students, compartment_percent),
        ("5", "Failed", failed_students, fail_percent),
        ("6", "A1 Grades", grade_count["A1"], round((grade_count["A1"]/total_grade_slots)*100,2)),
        ("7", "A2 Grades", grade_count["A2"], round((grade_count["A2"]/total_grade_slots)*100,2)),
        ("8", "B1 Grades", grade_count["B1"], round((grade_count["B1"]/total_grade_slots)*100,2)),
        ("9", "B2 Grades", grade_count["B2"], round((grade_count["B2"]/total_grade_slots)*100,2)),
        ("10", "C1 Grades", grade_count["C1"], round((grade_count["C1"]/total_grade_slots)*100,2)),
        ("11", "C2 Grades", grade_count["C2"], round((grade_count["C2"]/total_grade_slots)*100,2)),
        ("12", "D1 Grades", grade_count["D1"], round((grade_count["D1"]/total_grade_slots)*100,2)),
        ("13", "D2 Grades", grade_count["D2"], round((grade_count["D2"]/total_grade_slots)*100,2)),
        ("14", "E Grades", grade_count["E"], round((grade_count["E"]/total_grade_slots)*100,2)),
        ("15", "Highest % (Main 5)", maximum_marks if 'maximum_marks' in locals() else 0, highest_percent),
        ("16", "A1 in all subjects",str(A1_inall_sub),""),
        ("17", "School Average % (Main 5)", school_avg, ""),
        ("18", "School Score",school_score, ""),
    ]

    for row_data in rows_data:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        row[2].text = str(row_data[2])
        row[3].text = str(row_data[3])

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("Signature of PRINCIPAL".ljust(60) + "Signature of EXAM INCHARGE")

    # ========================= TEACHER PERFORMANCE SHEET =============================
    teacher_sheet = wb.create_sheet("Teacher Performance Detailed")

    teacher_sheet.append([
        "S.No", "Subject", "Teacher Name",
        "A1", "A1 %", "A2", "A2 %", "B1", "B1 %",
        "B2", "B2 %", "C1", "C1 %", "C2", "C2 %",
        "D1", "D1 %", "D2", "D2 %", "E", "E %",
        "Teacher Score"
    ])

    sno = 1
    subject_grade_distribution = {}
    for subject in sorted(all_subjects):

        # count grades for the subject
        grade_counts_sub = {
            "A1": 0, "A2": 0, "B1": 0, "B2": 0,
            "C1": 0, "C2": 0, "D1": 0, "D2": 0, "E": 0
        }

        total_students_subject = 0

        for roll, info in students.items():
            if subject in info["Grades"]:
                grade = info["Grades"][subject]
                if grade in grade_counts_sub:
                    grade_counts_sub[grade] += 1
                total_students_subject += 1

        subject_grade_distribution[subject] = grade_counts_sub

        if total_students_subject == 0:
            total_students_subject = 1

        grade_percent_sub = {
            g: round((grade_counts_sub[g] / total_students_subject) * 100, 2)
            for g in grade_counts_sub
        }

        teacher_name = teacher_map.get(subject, "")

        teacher_score = round(
            ((9*grade_counts_sub["A1"] +
              8*grade_counts_sub["A2"] +
              7*grade_counts_sub["B1"] +
              6*grade_counts_sub["B2"] +
              5*grade_counts_sub["C1"] +
              4*grade_counts_sub["C2"] +
              3*grade_counts_sub["D1"] +
              2*grade_counts_sub["D2"] +
              1*grade_counts_sub["E"]) / (9 * total_students_subject)) * 100, 2
        )

        teacher_sheet.append([
            sno,
            subject,
            teacher_name,
            grade_counts_sub["A1"], grade_percent_sub["A1"],
            grade_counts_sub["A2"], grade_percent_sub["A2"],
            grade_counts_sub["B1"], grade_percent_sub["B1"],
            grade_counts_sub["B2"], grade_percent_sub["B2"],
            grade_counts_sub["C1"], grade_percent_sub["C1"],
            grade_counts_sub["C2"], grade_percent_sub["C2"],
            grade_counts_sub["D1"], grade_percent_sub["D1"],
            grade_counts_sub["D2"], grade_percent_sub["D2"],
            grade_counts_sub["E"], grade_percent_sub["E"],
            teacher_score
        ])
        sno += 1

    # ================= SAVE INTO BYTESIO BUFFERS ====
    excel_buffer = io.BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    # ===== TOP 5 STUDENTS =====
    sorted_students = sorted(
        students.items(),
        key=lambda x: x[1]["Main5_Percent"],
        reverse=True
    )

    top5 = {}
    for roll, info in sorted_students[:5]:
        top5[info["Name"]] = info["Main5_Percent"]


    # ===== GRADE DISTRIBUTION =====
    grade_distribution = grade_count


    # ===== ANALYTICS OBJECT =====
    analytics = {
    "total_students": total_students,
    "present_students": present_students,
    "absent_students": absent_students,
    "result_later": rl_students,

    "pass_percent": pass_percent,
    "fail_percent": fail_percent,
    "compartment_percent": compartment_percent,

    "school_avg": school_avg,
    "highest_percent": highest_percent,

    "all_A1": A1_inall_sub,

    "grade_distribution": grade_count,

    "subject_grade_distribution": subject_grade_distribution,

    "top5": top5
    }


    return {
        "excel_file": excel_buffer,
        "word_file": word_buffer,
        "analytics": analytics
    }
